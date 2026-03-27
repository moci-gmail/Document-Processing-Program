import streamlit as st
import os
import locale
from docx import Document
from PyPDF2 import PdfReader
import re
import tempfile
import requests
import json
import pickle

# 多语言支持
# 翻译字典
translations = {
    "中文": {
        "page_title": "CleanDocs - AI文档规范化工具",
        "title": "CleanDocs - AI文档规范化自动化工具",
        "subheader": "一键排版、自动纠错、格式统一，告别手动整理文档的繁琐",
        "function_select": "选择功能",
        "functions": ["文档格式化", "标点&语法纠错", "批量文档处理", "AI语句润色", "标题层级生成", "引用格式规整", "自定义格式模板"],
        "language_settings": "语言设置",
        "interface_language": "界面语言",
        "processing_language": "处理语言",
        "processing_languages": ["自动检测", "中文", "繁體中文", "English"],
        "ai_config": "AI配置",
        "ai_model_select": "AI模型选择",
        "ai_models": ["本地Ollama", "OpenAI", "Anthropic Claude", "Google Gemini", "通义千问", "文心一言"],
        "ollama_url": "Ollama API地址",
        "model_name": "模型名称",
        "openai_api_key": "OpenAI API密钥",
        "anthropic_api_key": "Anthropic API密钥",
        "google_api_key": "Google API密钥",
        "api_key": "API密钥",
        "document_formatting": "文档格式化",
        "upload_doc": "上传Word或PDF文档，自动统一格式",
        "select_doc": "选择文档",
        "processing": "处理中...",
        "download_formatted": "下载格式化文档",
        "download_extracted": "下载提取的文本",
        "processing_complete": "文档处理完成！",
        "punctuation_correction": "标点&语法纠错",
        "input_option": "输入方式",
        "input_options": ["直接输入文本", "上传文档"],
        "enter_text": "输入需要纠错的文本",
        "start_correction": "开始纠错",
        "correction_result": "纠错结果",
        "download_corrected": "下载纠错文本",
        "download_corrected_doc": "下载纠错文档",
        "correction_complete": "纠错完成！",
        "document_correction_complete": "文档纠错完成！",
        "please_enter_text": "请输入文本",
        "batch_processing": "批量文档处理",
        "upload_multiple": "上传多个文档，批量完成格式规整",
        "select_multiple": "选择多个文档",
        "uploaded_count": "已上传 {count} 个文档",
        "start_batch": "开始批量处理",
        "batch_processing_": "批量处理中...",
        "download_processed": "下载处理后的 {file}",
        "download_extracted_": "下载提取的 {file}",
        "batch_complete": "批量处理完成！",
        "ai_polishing": "AI语句润色",
        "ai_polishing_desc": "针对课业、职场文案，优化语句通顺度，贴合学术/正式文风",
        "enter_polish_text": "输入需要润色的文本",
        "start_polishing": "开始润色",
        "ai_polishing_": "AI润色中...",
        "polishing_result": "润色结果",
        "download_polished": "下载润色文本",
        "download_polished_doc": "下载润色文档",
        "polishing_complete": "润色完成！",
        "document_polishing_complete": "文档润色完成！",
        "heading_generation": "标题层级自动生成",
        "heading_generation_desc": "自动识别文档标题、副标题，生成规范目录",
        "generated_toc": "生成的目录",
        "download_with_toc": "下载带目录的文档",
        "toc_generation_complete": "目录生成完成！",
        "no_headings_detected": "未识别到标题层级",
        "citation_formatting": "引用格式简易规整",
        "citation_formatting_desc": "针对essay、小论文，自动修正APA、MLA基础引用格式",
        "citation_style": "引用格式",
        "enter_citation_text": "输入需要规整引用格式的文本",
        "start_formatting": "开始规整",
        "formatting_result": "规整结果",
        "download_formatted_text": "下载规整文本",
        "download_formatted_doc": "下载规整文档",
        "citation_formatting_complete": "引用格式规整完成！",
        "custom_templates": "自定义格式模板",
        "custom_templates_desc": "保存专属排版格式，下次一键套用",
        "template_action": "操作",
        "template_actions": ["创建新模板", "应用现有模板"],
        "template_name": "模板名称",
        "font": "字体",
        "font_size": "字号",
        "line_spacing": "行距",
        "save_template": "保存模板",
        "template_saved": "模板 {name} 保存成功！",
        "please_enter_template_name": "请输入模板名称",
        "select_template": "选择模板",
        "apply_template": "应用模板",
        "applying_template": "应用模板中...",
        "download_template_applied": "下载应用模板后的文档",
        "template_applied": "模板应用完成！",
        "no_templates": "暂无保存的模板",
        "footer1": "© 2026 CleanDocs - 轻量级文档规范化工具",
        "footer2": "专为学生、自由撰稿人和职场人士设计",
        "ai_failed": "AI调用失败",
        "ollama_not_running": "本地Ollama未运行",
        "please_enter_api_key": "请输入{service} API密钥",
        "api_failed": "{service} API调用失败: {error}",
        "api_error": "{service} API调用出错: {error}",
        "feature_in_development": "{service}功能开发中",
        "please_select_valid_model": "请选择有效的AI模型"
    },
    "繁體中文": {
        "page_title": "CleanDocs - AI文檔規範化工具",
        "title": "CleanDocs - AI文檔規範化自動化工具",
        "subheader": "一鍵排版、自動糾錯、格式統一，告別手動整理文檔的繁瑣",
        "function_select": "選擇功能",
        "functions": ["文檔格式化", "標點&語法糾錯", "批量文檔處理", "AI語句潤色", "標題層級生成", "引用格式規整", "自定義格式模板"],
        "language_settings": "語言設定",
        "interface_language": "界面語言",
        "processing_language": "處理語言",
        "processing_languages": ["自動檢測", "中文", "繁體中文", "English"],
        "ai_config": "AI配置",
        "ai_model_select": "AI模型選擇",
        "ai_models": ["本地Ollama", "OpenAI", "Anthropic Claude", "Google Gemini", "通義千問", "文心一言"],
        "ollama_url": "Ollama API地址",
        "model_name": "模型名稱",
        "openai_api_key": "OpenAI API密鑰",
        "anthropic_api_key": "Anthropic API密鑰",
        "google_api_key": "Google API密鑰",
        "api_key": "API密鑰",
        "document_formatting": "文檔格式化",
        "upload_doc": "上傳Word或PDF文檔，自動統一格式",
        "select_doc": "選擇文檔",
        "processing": "處理中...",
        "download_formatted": "下載格式化文檔",
        "download_extracted": "下載提取的文本",
        "processing_complete": "文檔處理完成！",
        "punctuation_correction": "標點&語法糾錯",
        "input_option": "輸入方式",
        "input_options": ["直接輸入文本", "上傳文檔"],
        "enter_text": "輸入需要糾錯的文本",
        "start_correction": "開始糾錯",
        "correction_result": "糾錯結果",
        "download_corrected": "下載糾錯文本",
        "download_corrected_doc": "下載糾錯文檔",
        "correction_complete": "糾錯完成！",
        "document_correction_complete": "文檔糾錯完成！",
        "please_enter_text": "請輸入文本",
        "batch_processing": "批量文檔處理",
        "upload_multiple": "上傳多個文檔，批量完成格式規整",
        "select_multiple": "選擇多個文檔",
        "uploaded_count": "已上傳 {count} 個文檔",
        "start_batch": "開始批量處理",
        "batch_processing_": "批量處理中...",
        "download_processed": "下載處理後的 {file}",
        "download_extracted_": "下載提取的 {file}",
        "batch_complete": "批量處理完成！",
        "ai_polishing": "AI語句潤色",
        "ai_polishing_desc": "針對課業、職場文案，優化語句通順度，貼合學術/正式文風",
        "enter_polish_text": "輸入需要潤色的文本",
        "start_polishing": "開始潤色",
        "ai_polishing_": "AI潤色中...",
        "polishing_result": "潤色結果",
        "download_polished": "下載潤色文本",
        "download_polished_doc": "下載潤色文檔",
        "polishing_complete": "潤色完成！",
        "document_polishing_complete": "文檔潤色完成！",
        "heading_generation": "標題層級自動生成",
        "heading_generation_desc": "自動識別文檔標題、副標題，生成規範目錄",
        "generated_toc": "生成的目錄",
        "download_with_toc": "下載帶目錄的文檔",
        "toc_generation_complete": "目錄生成完成！",
        "no_headings_detected": "未識別到標題層級",
        "citation_formatting": "引用格式簡易規整",
        "citation_formatting_desc": "針對essay、小論文，自動修正APA、MLA基礎引用格式",
        "citation_style": "引用格式",
        "enter_citation_text": "輸入需要規整引用格式的文本",
        "start_formatting": "開始規整",
        "formatting_result": "規整結果",
        "download_formatted_text": "下載規整文本",
        "download_formatted_doc": "下載規整文檔",
        "citation_formatting_complete": "引用格式規整完成！",
        "custom_templates": "自定義格式模板",
        "custom_templates_desc": "保存專屬排版格式，下次一鍵套用",
        "template_action": "操作",
        "template_actions": ["建立新模板", "應用現有模板"],
        "template_name": "模板名稱",
        "font": "字體",
        "font_size": "字號",
        "line_spacing": "行距",
        "save_template": "保存模板",
        "template_saved": "模板 {name} 保存成功！",
        "please_enter_template_name": "請輸入模板名稱",
        "select_template": "選擇模板",
        "apply_template": "應用模板",
        "applying_template": "應用模板中...",
        "download_template_applied": "下載應用模板後的文檔",
        "template_applied": "模板應用完成！",
        "no_templates": "暫無保存的模板",
        "footer1": "© 2026 CleanDocs - 輕量級文檔規範化工具",
        "footer2": "專為學生、自由撰稿人和職場人士設計",
        "ai_failed": "AI調用失敗",
        "ollama_not_running": "本地Ollama未運行",
        "please_enter_api_key": "請輸入{service} API密鑰",
        "api_failed": "{service} API調用失敗: {error}",
        "api_error": "{service} API調用出錯: {error}",
        "feature_in_development": "{service}功能開發中",
        "please_select_valid_model": "請選擇有效的AI模型"
    },
    "English": {
        "page_title": "CleanDocs - AI Document Normalization Tool",
        "title": "CleanDocs - AI Document Normalization Automation Tool",
        "subheader": "One-click formatting, automatic error correction, format unification, say goodbye to the tedious manual document organization",
        "function_select": "Select Function",
        "functions": ["Document Formatting", "Punctuation & Grammar Correction", "Batch Document Processing", "AI Sentence Polishing", "Heading Hierarchy Generation", "Citation Formatting", "Custom Format Templates"],
        "language_settings": "Language Settings",
        "interface_language": "Interface Language",
        "processing_language": "Processing Language",
        "processing_languages": ["Auto Detect", "中文", "繁體中文", "English"],
        "ai_config": "AI Configuration",
        "ai_model_select": "AI Model Selection",
        "ai_models": ["Local Ollama", "OpenAI", "Anthropic Claude", "Google Gemini", "Tongyi Qianwen", "Wenxin Yiyan"],
        "ollama_url": "Ollama API Address",
        "model_name": "Model Name",
        "openai_api_key": "OpenAI API Key",
        "anthropic_api_key": "Anthropic API Key",
        "google_api_key": "Google API Key",
        "api_key": "API Key",
        "document_formatting": "Document Formatting",
        "upload_doc": "Upload Word or PDF document, automatically unify format",
        "select_doc": "Select Document",
        "processing": "Processing...",
        "download_formatted": "Download Formatted Document",
        "download_extracted": "Download Extracted Text",
        "processing_complete": "Document processing completed!",
        "punctuation_correction": "Punctuation & Grammar Correction",
        "input_option": "Input Method",
        "input_options": ["Direct Text Input", "Upload Document"],
        "enter_text": "Enter text for correction",
        "start_correction": "Start Correction",
        "correction_result": "Correction Result",
        "download_corrected": "Download Corrected Text",
        "download_corrected_doc": "Download Corrected Document",
        "correction_complete": "Correction completed!",
        "document_correction_complete": "Document correction completed!",
        "please_enter_text": "Please enter text",
        "batch_processing": "Batch Document Processing",
        "upload_multiple": "Upload multiple documents, batch complete format normalization",
        "select_multiple": "Select Multiple Documents",
        "uploaded_count": "Uploaded {count} documents",
        "start_batch": "Start Batch Processing",
        "batch_processing_": "Batch processing...",
        "download_processed": "Download processed {file}",
        "download_extracted_": "Download extracted {file}",
        "batch_complete": "Batch processing completed!",
        "ai_polishing": "AI Sentence Polishing",
        "ai_polishing_desc": "For academic and workplace documents, optimize sentence fluency, fit academic/formal style",
        "enter_polish_text": "Enter text for polishing",
        "start_polishing": "Start Polishing",
        "ai_polishing_": "AI polishing...",
        "polishing_result": "Polishing Result",
        "download_polished": "Download Polished Text",
        "download_polished_doc": "Download Polished Document",
        "polishing_complete": "Polishing completed!",
        "document_polishing_complete": "Document polishing completed!",
        "heading_generation": "Heading Hierarchy Generation",
        "heading_generation_desc": "Automatically identify document headings and subheadings, generate standardized table of contents",
        "generated_toc": "Generated Table of Contents",
        "download_with_toc": "Download Document with TOC",
        "toc_generation_complete": "Table of contents generation completed!",
        "no_headings_detected": "No headings detected",
        "citation_formatting": "Citation Formatting",
        "citation_formatting_desc": "For essays and papers, automatically correct basic APA, MLA citation formats",
        "citation_style": "Citation Style",
        "enter_citation_text": "Enter text for citation formatting",
        "start_formatting": "Start Formatting",
        "formatting_result": "Formatting Result",
        "download_formatted_text": "Download Formatted Text",
        "download_formatted_doc": "Download Formatted Document",
        "citation_formatting_complete": "Citation formatting completed!",
        "custom_templates": "Custom Format Templates",
        "custom_templates_desc": "Save exclusive formatting styles, apply with one click next time",
        "template_action": "Action",
        "template_actions": ["Create New Template", "Apply Existing Template"],
        "template_name": "Template Name",
        "font": "Font",
        "font_size": "Font Size",
        "line_spacing": "Line Spacing",
        "save_template": "Save Template",
        "template_saved": "Template {name} saved successfully!",
        "please_enter_template_name": "Please enter template name",
        "select_template": "Select Template",
        "apply_template": "Apply Template",
        "applying_template": "Applying template...",
        "download_template_applied": "Download Document with Applied Template",
        "template_applied": "Template applied successfully!",
        "no_templates": "No saved templates",
        "footer1": "© 2026 CleanDocs - Lightweight Document Normalization Tool",
        "footer2": "Designed for students, freelance writers, and professionals",
        "ai_failed": "AI call failed",
        "ollama_not_running": "Local Ollama not running",
        "please_enter_api_key": "Please enter {service} API key",
        "api_failed": "{service} API call failed: {error}",
        "api_error": "{service} API call error: {error}",
        "feature_in_development": "{service} feature in development",
        "please_select_valid_model": "Please select a valid AI model"
    }
}

# 获取翻译
def get_translation(key, lang, **kwargs):
    try:
        text = translations[lang][key]
        if kwargs:
            return text.format(**kwargs)
        return text
    except KeyError:
        return key

# 检测系统语言
try:
    system_lang = locale.getdefaultlocale()[0]
    if system_lang.startswith('zh_TW') or system_lang.startswith('zh_HK'):
        default_language = "繁體中文"
    elif system_lang.startswith('zh'):
        default_language = "中文"
    else:
        default_language = "English"
except:
    default_language = "中文"

# 语言选择
language = st.sidebar.selectbox(
    get_translation("interface_language", "中文"),
    ["中文", "繁體中文", "English"],
    index=["中文", "繁體中文", "English"].index(default_language)
)

# 设置页面配置
st.set_page_config(
    page_title=get_translation("page_title", language),
    page_icon="📄",
    layout="wide"
)

# 自定义CSS来隐藏Streamlit内置的Deploy按钮和订阅功能
st.markdown("""
<style>
/* 隐藏汉堡菜单中的Deploy按钮 */
[data-testid="stSidebarNav"] ul li:nth-child(4) {
    display: none !important;
}

/* 隐藏订阅功能 */
[data-testid="stSidebarNav"] ul li:nth-child(5) {
    display: none !important;
}

/* 隐藏底部的Deploy相关文本 */
footer {
    display: none !important;
}
</style>
""", unsafe_allow_html=True)

# 标题和描述
st.title(get_translation("title", language))
st.subheader(get_translation("subheader", language))

# 侧边栏
st.sidebar.header(get_translation("function_select", language))
function_option = st.sidebar.selectbox(
    get_translation("function_select", language),
    get_translation("functions", language)
)

# 语言选择
st.sidebar.header(get_translation("language_settings", language))

# 处理语言选择
processing_language = st.sidebar.selectbox(
    get_translation("processing_language", language),
    get_translation("processing_languages", language)
)

# AI配置
st.sidebar.header(get_translation("ai_config", language))
ai_model = st.sidebar.selectbox(
    get_translation("ai_model_select", language),
    get_translation("ai_models", language)
)

if ai_model == get_translation("ai_models", language)[0]:  # 本地Ollama
    ollama_url = st.sidebar.text_input(get_translation("ollama_url", language), "http://localhost:11434")
    ollama_model = st.sidebar.text_input(get_translation("model_name", language), "llama3")
elif ai_model == get_translation("ai_models", language)[1]:  # OpenAI
    api_key = st.sidebar.text_input(get_translation("openai_api_key", language), type="password")
    openai_model = st.sidebar.selectbox(get_translation("model_name", language), ["gpt-4o", "gpt-4-turbo", "gpt-3.5-turbo"])
elif ai_model == get_translation("ai_models", language)[2]:  # Anthropic Claude
    api_key = st.sidebar.text_input(get_translation("anthropic_api_key", language), type="password")
    claude_model = st.sidebar.selectbox(get_translation("model_name", language), ["claude-3-opus-20240229", "claude-3-sonnet-20240229"])
elif ai_model == get_translation("ai_models", language)[3]:  # Google Gemini
    api_key = st.sidebar.text_input(get_translation("google_api_key", language), type="password")
    gemini_model = st.sidebar.selectbox(get_translation("model_name", language), ["gemini-1.5-pro", "gemini-1.5-flash"])
elif ai_model == get_translation("ai_models", language)[4]:  # 通义千问
    api_key = st.sidebar.text_input(get_translation("api_key", language), type="password")
elif ai_model == get_translation("ai_models", language)[5]:  # 文心一言
    api_key = st.sidebar.text_input(get_translation("api_key", language), type="password")

# 自定义格式模板管理
template_dir = "templates"
if not os.path.exists(template_dir):
    os.makedirs(template_dir)

# 主功能实现
if function_option == get_translation("functions", language)[0]:  # 文档格式化
    st.header(get_translation("document_formatting", language))
    st.write(get_translation("upload_doc", language))
    
    uploaded_file = st.file_uploader(get_translation("select_doc", language), type=["docx", "pdf"])
    
    if uploaded_file is not None:
        with st.spinner(get_translation("processing", language)):
            # 处理Word文档
            if uploaded_file.name.endswith(".docx"):
                doc = Document(uploaded_file)
                
                # 统一格式
                for paragraph in doc.paragraphs:
                    # 设置字体和字号
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = doc.styles['Normal'].font.size
                
                # 保存处理后的文档
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
                
                # 提供下载
                with open(tmp_path, "rb") as f:
                    st.download_button(
                        label=get_translation("download_formatted", language),
                        data=f,
                        file_name="formatted_" + uploaded_file.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # 清理临时文件
                os.unlink(tmp_path)
                
            # 处理PDF文档
            elif uploaded_file.name.endswith(".pdf"):
                reader = PdfReader(uploaded_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                
                # 保存为Word文档
                doc = Document()
                doc.add_paragraph(text)
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
                
                # 提供下载
                with open(tmp_path, "rb") as f:
                    st.download_button(
                        label=get_translation("download_extracted", language),
                        data=f,
                        file_name="extracted_" + uploaded_file.name.replace(".pdf", ".docx"),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                # 清理临时文件
                os.unlink(tmp_path)
                
            st.success(get_translation("processing_complete", language))

elif function_option == get_translation("functions", language)[1]:  # 标点&语法纠错
    st.header(get_translation("punctuation_correction", language))
    st.write(get_translation("upload_doc", language))
    
    # 文本输入选项
    input_option = st.radio(get_translation("input_option", language), get_translation("input_options", language))
    
    if input_option == get_translation("input_options", language)[0]:  # 直接输入文本
        text = st.text_area(get_translation("enter_text", language), height=300)
        
        if st.button(get_translation("start_correction", language)):
            if text:
                with st.spinner(get_translation("processing", language)):
                    # 基本纠错处理
                    # 修正中英文标点
                    text = text.replace("，", ", ").replace("。", ". ").replace("！", "! ").replace("？", "? ")
                    # 修正多余空格
                    text = re.sub(r'\s+', ' ', text)
                    # 修正大小写
                    sentences = re.split(r'[.!?]+', text)
                    corrected_text = ""
                    for sentence in sentences:
                        if sentence.strip():
                            corrected_text += sentence.strip().capitalize() + ". "
                    
                    st.subheader(get_translation("correction_result", language))
                    st.write(corrected_text)
                    
                    # 提供下载
                    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
                        tmp.write(corrected_text.encode('utf-8'))
                        tmp_path = tmp.name
                    
                    with open(tmp_path, "rb") as f:
                        st.download_button(
                            label=get_translation("download_corrected", language),
                            data=f,
                            file_name="corrected_text.txt",
                            mime="text/plain"
                        )
                    
                    os.unlink(tmp_path)
                    st.success(get_translation("correction_complete", language))
            else:
                st.warning(get_translation("please_enter_text", language))
    
    else:  # 上传文档
        uploaded_file = st.file_uploader(get_translation("select_doc", language), type=["docx", "txt"])
        
        if uploaded_file is not None:
            with st.spinner(get_translation("processing", language)):
                # 读取文档内容
                if uploaded_file.name.endswith(".docx"):
                    doc = Document(uploaded_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = uploaded_file.getvalue().decode('utf-8')
                
                # 基本纠错处理
                text = text.replace("，", ", ").replace("。", ". ").replace("！", "! ").replace("？", "? ")
                text = re.sub(r'\s+', ' ', text)
                sentences = re.split(r'[.!?]+', text)
                corrected_text = ""
                for sentence in sentences:
                    if sentence.strip():
                        corrected_text += sentence.strip().capitalize() + ". "
                
                # 保存为Word文档
                doc = Document()
                doc.add_paragraph(corrected_text)
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
                
                # 提供下载
                with open(tmp_path, "rb") as f:
                    st.download_button(
                        label=get_translation("download_corrected_doc", language),
                        data=f,
                        file_name="corrected_" + uploaded_file.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                os.unlink(tmp_path)
                st.success(get_translation("document_correction_complete", language))

elif function_option == get_translation("functions", language)[2]:  # 批量文档处理
    st.header(get_translation("batch_processing", language))
    st.write(get_translation("upload_multiple", language))
    
    uploaded_files = st.file_uploader(get_translation("select_multiple", language), type=["docx", "pdf"], accept_multiple_files=True)
    
    if uploaded_files:
        st.write(get_translation("uploaded_count", language, count=len(uploaded_files)))
        
        if st.button(get_translation("start_batch", language)):
            with st.spinner(get_translation("batch_processing_", language)):
                # 处理每个文档
                for file in uploaded_files:
                    if file.name.endswith(".docx"):
                        doc = Document(file)
                        # 统一格式
                        for paragraph in doc.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = "Calibri"
                                run.font.size = doc.styles['Normal'].font.size
                        
                        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                            doc.save(tmp.name)
                            tmp_path = tmp.name
                        
                        with open(tmp_path, "rb") as f:
                            st.download_button(
                                label=get_translation("download_processed", language, file=file.name),
                                data=f,
                                file_name="formatted_" + file.name,
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        os.unlink(tmp_path)
                    
                    elif file.name.endswith(".pdf"):
                        reader = PdfReader(file)
                        text = ""
                        for page in reader.pages:
                            text += page.extract_text() + "\n"
                        
                        doc = Document()
                        doc.add_paragraph(text)
                        
                        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                            doc.save(tmp.name)
                            tmp_path = tmp.name
                        
                        with open(tmp_path, "rb") as f:
                            st.download_button(
                                label=get_translation("download_extracted_", language, file=file.name),
                                data=f,
                                file_name="extracted_" + file.name.replace(".pdf", ".docx"),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        os.unlink(tmp_path)
                
                st.success(get_translation("batch_complete", language))

elif function_option == get_translation("functions", language)[3]:  # AI语句润色
    st.header(get_translation("ai_polishing", language))
    st.write(get_translation("ai_polishing_desc", language))
    
    # 文本输入选项
    input_option = st.radio(get_translation("input_option", language), get_translation("input_options", language))
    
    if input_option == get_translation("input_options", language)[0]:  # 直接输入文本
        text = st.text_area(get_translation("enter_polish_text", language), height=300)
        
        if st.button(get_translation("start_polishing", language)):
            if text:
                with st.spinner(get_translation("ai_polishing_", language)):
                    # 调用AI进行润色
                    def call_ai(prompt):
                        if ai_model == get_translation("ai_models", language)[0]:  # 本地Ollama
                            # 调用本地Ollama
                            try:
                                response = requests.post(
                                    f"{ollama_url}/api/generate",
                                    json={
                                        "model": ollama_model,
                                        "prompt": prompt,
                                        "stream": False
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json().get("response", "")
                                else:
                                    return get_translation("ai_failed", language)
                            except:
                                return get_translation("ollama_not_running", language)
                        elif ai_model == get_translation("ai_models", language)[1]:  # OpenAI
                            # 调用OpenAI API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="OpenAI")
                            try:
                                response = requests.post(
                                    "https://api.openai.com/v1/chat/completions",
                                    headers={
                                        "Authorization": f"Bearer {api_key}",
                                        "Content-Type": "application/json"
                                    },
                                    json={
                                        "model": openai_model,
                                        "messages": [
                                            {"role": "user", "content": prompt}
                                        ],
                                        "temperature": 0.3
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json()["choices"][0]["message"]["content"]
                                else:
                                    return get_translation("api_failed", language, service="OpenAI", error=response.text)
                            except Exception as e:
                                return get_translation("api_error", language, service="OpenAI", error=str(e))
                        elif ai_model == get_translation("ai_models", language)[2]:  # Anthropic Claude
                            # 调用Anthropic API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="Anthropic")
                            try:
                                response = requests.post(
                                    "https://api.anthropic.com/v1/messages",
                                    headers={
                                        "x-api-key": api_key,
                                        "Content-Type": "application/json",
                                        "anthropic-version": "2023-06-01"
                                    },
                                    json={
                                        "model": claude_model,
                                        "messages": [
                                            {"role": "user", "content": prompt}
                                        ],
                                        "temperature": 0.3
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json()["content"][0]["text"]
                                else:
                                    return get_translation("api_failed", language, service="Anthropic", error=response.text)
                            except Exception as e:
                                return get_translation("api_error", language, service="Anthropic", error=str(e))
                        elif ai_model == get_translation("ai_models", language)[3]:  # Google Gemini
                            # 调用Google Gemini API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="Google")
                            try:
                                response = requests.post(
                                    f"https://generativelanguage.googleapis.com/v1/models/{gemini_model}:generateContent?key={api_key}",
                                    headers={
                                        "Content-Type": "application/json"
                                    },
                                    json={
                                        "contents": [
                                            {"parts": [{"text": prompt}]}
                                        ],
                                        "generationConfig": {
                                            "temperature": 0.3
                                        }
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json()["candidates"][0]["content"]["parts"][0]["text"]
                                else:
                                    return get_translation("api_failed", language, service="Google Gemini", error=response.text)
                            except Exception as e:
                                return get_translation("api_error", language, service="Google Gemini", error=str(e))
                        elif ai_model == get_translation("ai_models", language)[4]:  # 通义千问
                            # 调用通义千问API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="")
                            return get_translation("feature_in_development", language, service="通义千问")
                        elif ai_model == get_translation("ai_models", language)[5]:  # 文心一言
                            # 调用文心一言API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="")
                            return get_translation("feature_in_development", language, service="文心一言")
                        else:
                            return get_translation("please_select_valid_model", language)
                    
                    prompt = f"请润色以下文本，保持原意，优化语句通顺度，贴合学术/正式文风，避免AI痕迹：\n\n{text}"
                    polished_text = call_ai(prompt)
                    
                    st.subheader(get_translation("polishing_result", language))
                    st.write(polished_text)
                    
                    # 提供下载
                    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
                        tmp.write(polished_text.encode('utf-8'))
                        tmp_path = tmp.name
                    
                    with open(tmp_path, "rb") as f:
                        st.download_button(
                            label=get_translation("download_polished", language),
                            data=f,
                            file_name="polished_text.txt",
                            mime="text/plain"
                        )
                    
                    os.unlink(tmp_path)
                    st.success(get_translation("polishing_complete", language))
            else:
                st.warning(get_translation("please_enter_text", language))
    
    else:  # 上传文档
        uploaded_file = st.file_uploader(get_translation("select_doc", language), type=["docx", "txt"])
        
        if uploaded_file is not None:
            with st.spinner(get_translation("processing", language)):
                # 读取文档内容
                if uploaded_file.name.endswith(".docx"):
                    doc = Document(uploaded_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = uploaded_file.getvalue().decode('utf-8')
                
                # 调用AI进行润色
                def call_ai(prompt):
                    if ai_model == get_translation("ai_models", language)[0]:  # 本地Ollama
                        try:
                            response = requests.post(
                                f"{ollama_url}/api/generate",
                                json={
                                    "model": ollama_model,
                                    "prompt": prompt,
                                    "stream": False
                                }
                            )
                            if response.status_code == 200:
                                return response.json().get("response", "")
                            else:
                                return get_translation("ai_failed", language)
                        except:
                            return get_translation("ollama_not_running", language)
                    elif ai_model == get_translation("ai_models", language)[1]:  # OpenAI
                        # 调用OpenAI API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="OpenAI")
                        try:
                            response = requests.post(
                                "https://api.openai.com/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {api_key}",
                                    "Content-Type": "application/json"
                                },
                                json={
                                    "model": openai_model,
                                    "messages": [
                                        {"role": "user", "content": prompt}
                                    ],
                                    "temperature": 0.3
                                }
                            )
                            if response.status_code == 200:
                                return response.json()["choices"][0]["message"]["content"]
                            else:
                                return get_translation("api_failed", language, service="OpenAI", error=response.text)
                        except Exception as e:
                            return get_translation("api_error", language, service="OpenAI", error=str(e))
                    elif ai_model == get_translation("ai_models", language)[2]:  # Anthropic Claude
                        # 调用Anthropic API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="Anthropic")
                        try:
                            response = requests.post(
                                "https://api.anthropic.com/v1/messages",
                                headers={
                                    "x-api-key": api_key,
                                    "Content-Type": "application/json",
                                    "anthropic-version": "2023-06-01"
                                },
                                json={
                                    "model": claude_model,
                                    "messages": [
                                        {"role": "user", "content": prompt}
                                    ],
                                    "temperature": 0.3
                                }
                            )
                            if response.status_code == 200:
                                return response.json()["content"][0]["text"]
                            else:
                                return get_translation("api_failed", language, service="Anthropic", error=response.text)
                        except Exception as e:
                            return get_translation("api_error", language, service="Anthropic", error=str(e))
                    elif ai_model == get_translation("ai_models", language)[3]:  # Google Gemini
                        # 调用Google Gemini API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="Google")
                        try:
                            response = requests.post(
                                f"https://generativelanguage.googleapis.com/v1/models/{gemini_model}:generateContent?key={api_key}",
                                headers={
                                    "Content-Type": "application/json"
                                },
                                json={
                                    "contents": [
                                        {"parts": [{"text": prompt}]}
                                    ],
                                    "generationConfig": {
                                        "temperature": 0.3
                                    }
                                }
                            )
                            if response.status_code == 200:
                                return response.json()["candidates"][0]["content"]["parts"][0]["text"]
                            else:
                                return get_translation("api_failed", language, service="Google Gemini", error=response.text)
                        except Exception as e:
                            return get_translation("api_error", language, service="Google Gemini", error=str(e))
                    elif ai_model == get_translation("ai_models", language)[4]:  # 通义千问
                        # 调用通义千问API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="")
                        return get_translation("feature_in_development", language, service="通义千问")
                    elif ai_model == get_translation("ai_models", language)[5]:  # 文心一言
                        # 调用文心一言API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="")
                        return get_translation("feature_in_development", language, service="文心一言")
                    else:
                        return get_translation("please_select_valid_model", language)
                
                prompt = f"请润色以下文本，保持原意，优化语句通顺度，贴合学术/正式文风，避免AI痕迹：\n\n{text}"
                polished_text = call_ai(prompt)
                
                # 保存为Word文档
                doc = Document()
                doc.add_paragraph(polished_text)
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
                
                # 提供下载
                with open(tmp_path, "rb") as f:
                    st.download_button(
                        label=get_translation("download_polished_doc", language),
                        data=f,
                        file_name="polished_" + uploaded_file.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                os.unlink(tmp_path)
                st.success(get_translation("document_polishing_complete", language))

elif function_option == get_translation("functions", language)[4]:  # 标题层级生成
    st.header(get_translation("heading_generation", language))
    st.write(get_translation("heading_generation_desc", language))
    
    uploaded_file = st.file_uploader(get_translation("select_doc", language), type=["docx"])
    
    if uploaded_file is not None:
        with st.spinner(get_translation("processing", language)):
            doc = Document(uploaded_file)
            
            # 识别标题层级
            headings = []
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.replace('Heading ', ''))
                    headings.append((level, paragraph.text, i))
            
            # 生成目录
            if headings:
                st.subheader(get_translation("generated_toc", language))
                toc = []
                for level, text, index in headings:
                    indent = "  " * (level - 1)
                    toc.append(f"{indent}{level}. {text}")
                
                for item in toc:
                    st.write(item)
                
                # 在文档开头添加目录
                new_doc = Document()
                new_doc.add_heading(get_translation("generated_toc", language), level=0)
                for item in toc:
                    new_doc.add_paragraph(item)
                new_doc.add_page_break()
                
                # 添加原文档内容
                for paragraph in doc.paragraphs:
                    new_paragraph = new_doc.add_paragraph()
                    new_paragraph.text = paragraph.text
                    new_paragraph.style = paragraph.style
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    new_doc.save(tmp.name)
                    tmp_path = tmp.name
                
                with open(tmp_path, "rb") as f:
                    st.download_button(
                        label=get_translation("download_with_toc", language),
                        data=f,
                        file_name="with_toc_" + uploaded_file.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                os.unlink(tmp_path)
                st.success(get_translation("toc_generation_complete", language))
            else:
                st.warning(get_translation("no_headings_detected", language))

elif function_option == get_translation("functions", language)[5]:  # 引用格式规整
    st.header(get_translation("citation_formatting", language))
    st.write(get_translation("citation_formatting_desc", language))
    
    # 文本输入选项
    input_option = st.radio(get_translation("input_option", language), get_translation("input_options", language))
    citation_style = st.selectbox(get_translation("citation_style", language), ["APA", "MLA"])
    
    if input_option == get_translation("input_options", language)[0]:  # 直接输入文本
        text = st.text_area(get_translation("enter_citation_text", language), height=300)
        
        if st.button(get_translation("start_formatting", language)):
            if text:
                with st.spinner(get_translation("processing", language)):
                    # 这里可以添加引用格式规整的逻辑
                    # 简单示例：识别可能的引用并进行格式化
                    # 实际实现需要更复杂的正则表达式和规则
                    
                    # 调用AI进行引用格式规整
                    def call_ai(prompt):
                        if ai_model == get_translation("ai_models", language)[0]:  # 本地Ollama
                            try:
                                response = requests.post(
                                    f"{ollama_url}/api/generate",
                                    json={
                                        "model": ollama_model,
                                        "prompt": prompt,
                                        "stream": False
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json().get("response", "")
                                else:
                                    return get_translation("ai_failed", language)
                            except:
                                return get_translation("ollama_not_running", language)
                        elif ai_model == get_translation("ai_models", language)[1]:  # OpenAI
                            # 调用OpenAI API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="OpenAI")
                            try:
                                response = requests.post(
                                    "https://api.openai.com/v1/chat/completions",
                                    headers={
                                        "Authorization": f"Bearer {api_key}",
                                        "Content-Type": "application/json"
                                    },
                                    json={
                                        "model": openai_model,
                                        "messages": [
                                            {"role": "user", "content": prompt}
                                        ],
                                        "temperature": 0.3
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json()["choices"][0]["message"]["content"]
                                else:
                                    return get_translation("api_failed", language, service="OpenAI", error=response.text)
                            except Exception as e:
                                return get_translation("api_error", language, service="OpenAI", error=str(e))
                        elif ai_model == get_translation("ai_models", language)[2]:  # Anthropic Claude
                            # 调用Anthropic API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="Anthropic")
                            try:
                                response = requests.post(
                                    "https://api.anthropic.com/v1/messages",
                                    headers={
                                        "x-api-key": api_key,
                                        "Content-Type": "application/json",
                                        "anthropic-version": "2023-06-01"
                                    },
                                    json={
                                        "model": claude_model,
                                        "messages": [
                                            {"role": "user", "content": prompt}
                                        ],
                                        "temperature": 0.3
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json()["content"][0]["text"]
                                else:
                                    return get_translation("api_failed", language, service="Anthropic", error=response.text)
                            except Exception as e:
                                return get_translation("api_error", language, service="Anthropic", error=str(e))
                        elif ai_model == get_translation("ai_models", language)[3]:  # Google Gemini
                            # 调用Google Gemini API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="Google")
                            try:
                                response = requests.post(
                                    f"https://generativelanguage.googleapis.com/v1/models/{gemini_model}:generateContent?key={api_key}",
                                    headers={
                                        "Content-Type": "application/json"
                                    },
                                    json={
                                        "contents": [
                                            {"parts": [{"text": prompt}]}
                                        ],
                                        "generationConfig": {
                                            "temperature": 0.3
                                        }
                                    }
                                )
                                if response.status_code == 200:
                                    return response.json()["candidates"][0]["content"]["parts"][0]["text"]
                                else:
                                    return get_translation("api_failed", language, service="Google Gemini", error=response.text)
                            except Exception as e:
                                return get_translation("api_error", language, service="Google Gemini", error=str(e))
                        elif ai_model == get_translation("ai_models", language)[4]:  # 通义千问
                            # 调用通义千问API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="")
                            return get_translation("feature_in_development", language, service="通义千问")
                        elif ai_model == get_translation("ai_models", language)[5]:  # 文心一言
                            # 调用文心一言API
                            if not api_key:
                                return get_translation("please_enter_api_key", language, service="")
                            return get_translation("feature_in_development", language, service="文心一言")
                        else:
                            return get_translation("please_select_valid_model", language)
                    
                    prompt = f"请将以下文本中的引用格式修正为{citation_style}格式：\n\n{text}"
                    formatted_text = call_ai(prompt)
                    
                    st.subheader(get_translation("formatting_result", language))
                    st.write(formatted_text)
                    
                    # 提供下载
                    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as tmp:
                        tmp.write(formatted_text.encode('utf-8'))
                        tmp_path = tmp.name
                    
                    with open(tmp_path, "rb") as f:
                        st.download_button(
                            label=get_translation("download_formatted_text", language),
                            data=f,
                            file_name="formatted_citations.txt",
                            mime="text/plain"
                        )
                    
                    os.unlink(tmp_path)
                    st.success(get_translation("citation_formatting_complete", language))
            else:
                st.warning(get_translation("please_enter_text", language))
    
    else:  # 上传文档
        uploaded_file = st.file_uploader(get_translation("select_doc", language), type=["docx", "txt"])
        
        if uploaded_file is not None:
            with st.spinner(get_translation("processing", language)):
                # 读取文档内容
                if uploaded_file.name.endswith(".docx"):
                    doc = Document(uploaded_file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = uploaded_file.getvalue().decode('utf-8')
                
                # 调用AI进行引用格式规整
                def call_ai(prompt):
                    if ai_model == get_translation("ai_models", language)[0]:  # 本地Ollama
                        try:
                            response = requests.post(
                                f"{ollama_url}/api/generate",
                                json={
                                    "model": ollama_model,
                                    "prompt": prompt,
                                    "stream": False
                                }
                            )
                            if response.status_code == 200:
                                return response.json().get("response", "")
                            else:
                                return get_translation("ai_failed", language)
                        except:
                            return get_translation("ollama_not_running", language)
                    elif ai_model == get_translation("ai_models", language)[1]:  # OpenAI
                        # 调用OpenAI API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="OpenAI")
                        try:
                            response = requests.post(
                                "https://api.openai.com/v1/chat/completions",
                                headers={
                                    "Authorization": f"Bearer {api_key}",
                                    "Content-Type": "application/json"
                                },
                                json={
                                    "model": openai_model,
                                    "messages": [
                                        {"role": "user", "content": prompt}
                                    ],
                                    "temperature": 0.3
                                }
                            )
                            if response.status_code == 200:
                                return response.json()["choices"][0]["message"]["content"]
                            else:
                                return get_translation("api_failed", language, service="OpenAI", error=response.text)
                        except Exception as e:
                            return get_translation("api_error", language, service="OpenAI", error=str(e))
                    elif ai_model == get_translation("ai_models", language)[2]:  # Anthropic Claude
                        # 调用Anthropic API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="Anthropic")
                        try:
                            response = requests.post(
                                "https://api.anthropic.com/v1/messages",
                                headers={
                                    "x-api-key": api_key,
                                    "Content-Type": "application/json",
                                    "anthropic-version": "2023-06-01"
                                },
                                json={
                                    "model": claude_model,
                                    "messages": [
                                        {"role": "user", "content": prompt}
                                    ],
                                    "temperature": 0.3
                                }
                            )
                            if response.status_code == 200:
                                return response.json()["content"][0]["text"]
                            else:
                                return get_translation("api_failed", language, service="Anthropic", error=response.text)
                        except Exception as e:
                            return get_translation("api_error", language, service="Anthropic", error=str(e))
                    elif ai_model == get_translation("ai_models", language)[3]:  # Google Gemini
                        # 调用Google Gemini API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="Google")
                        try:
                            response = requests.post(
                                f"https://generativelanguage.googleapis.com/v1/models/{gemini_model}:generateContent?key={api_key}",
                                headers={
                                    "Content-Type": "application/json"
                                },
                                json={
                                    "contents": [
                                        {"parts": [{"text": prompt}]}
                                    ],
                                    "generationConfig": {
                                        "temperature": 0.3
                                    }
                                }
                            )
                            if response.status_code == 200:
                                return response.json()["candidates"][0]["content"]["parts"][0]["text"]
                            else:
                                return get_translation("api_failed", language, service="Google Gemini", error=response.text)
                        except Exception as e:
                            return get_translation("api_error", language, service="Google Gemini", error=str(e))
                    elif ai_model == get_translation("ai_models", language)[4]:  # 通义千问
                        # 调用通义千问API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="")
                        return get_translation("feature_in_development", language, service="通义千问")
                    elif ai_model == get_translation("ai_models", language)[5]:  # 文心一言
                        # 调用文心一言API
                        if not api_key:
                            return get_translation("please_enter_api_key", language, service="")
                        return get_translation("feature_in_development", language, service="文心一言")
                    else:
                        return get_translation("please_select_valid_model", language)
                
                prompt = f"请将以下文本中的引用格式修正为{citation_style}格式：\n\n{text}"
                formatted_text = call_ai(prompt)
                
                # 保存为Word文档
                doc = Document()
                doc.add_paragraph(formatted_text)
                
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    doc.save(tmp.name)
                    tmp_path = tmp.name
                
                # 提供下载
                with open(tmp_path, "rb") as f:
                    st.download_button(
                        label=get_translation("download_formatted_doc", language),
                        data=f,
                        file_name="formatted_citations_" + uploaded_file.name,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
                os.unlink(tmp_path)
                st.success(get_translation("citation_formatting_complete", language))

elif function_option == get_translation("functions", language)[6]:  # 自定义格式模板
    st.header(get_translation("custom_templates", language))
    st.write(get_translation("custom_templates_desc", language))
    
    # 模板操作选项
    template_action = st.radio(get_translation("template_action", language), get_translation("template_actions", language))
    
    if template_action == get_translation("template_actions", language)[0]:  # 创建新模板
        template_name = st.text_input(get_translation("template_name", language))
        font_name = st.text_input(get_translation("font", language), "Calibri")
        font_size = st.number_input(get_translation("font_size", language), min_value=8, max_value=24, value=12)
        line_spacing = st.number_input(get_translation("line_spacing", language), min_value=1.0, max_value=3.0, value=1.5)
        
        if st.button(get_translation("save_template", language)):
            if template_name:
                template_data = {
                    "font_name": font_name,
                    "font_size": font_size,
                    "line_spacing": line_spacing
                }
                
                with open(os.path.join(template_dir, f"{template_name}.pkl"), "wb") as f:
                    pickle.dump(template_data, f)
                
                st.success(get_translation("template_saved", language, name=template_name))
            else:
                st.warning(get_translation("please_enter_template_name", language))
    
    else:  # 应用现有模板
        # 列出现有模板
        templates = [f[:-4] for f in os.listdir(template_dir) if f.endswith(".pkl")]
        if templates:
            selected_template = st.selectbox(get_translation("select_template", language), templates)
            uploaded_file = st.file_uploader(get_translation("select_doc", language), type=["docx"])
            
            if uploaded_file is not None:
                if st.button(get_translation("apply_template", language)):
                    with st.spinner(get_translation("applying_template", language)):
                        # 加载模板
                        with open(os.path.join(template_dir, f"{selected_template}.pkl"), "rb") as f:
                            template_data = pickle.load(f)
                        
                        # 应用模板到文档
                        doc = Document(uploaded_file)
                        
                        # 设置字体和字号
                        for paragraph in doc.paragraphs:
                            for run in paragraph.runs:
                                run.font.name = template_data["font_name"]
                                run.font.size = doc.styles['Normal'].font.size
                        
                        # 这里可以添加更多格式设置，如行距等
                        
                        # 保存处理后的文档
                        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                            doc.save(tmp.name)
                            tmp_path = tmp.name
                        
                        # 提供下载
                        with open(tmp_path, "rb") as f:
                            st.download_button(
                                label=get_translation("download_template_applied", language),
                                data=f,
                                file_name=f"template_{selected_template}_{uploaded_file.name}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        os.unlink(tmp_path)
                        st.success(get_translation("template_applied", language))
        else:
            st.warning(get_translation("no_templates", language))

# 页脚
st.sidebar.markdown("---")
st.sidebar.markdown(get_translation("footer1", language))
st.sidebar.markdown(get_translation("footer2", language))